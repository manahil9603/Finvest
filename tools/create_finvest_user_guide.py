from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "finvest_user_guide_assets"
DOCX_PATH = OUT_DIR / "Finvest_Website_User_Guide.docx"

ACCENT = RGBColor(124, 58, 237)
GREEN = RGBColor(16, 185, 129)
BLUE = RGBColor(59, 130, 246)
AMBER = RGBColor(245, 158, 11)
DARK = RGBColor(24, 24, 27)
MUTED = RGBColor(82, 82, 91)


def set_cell_shading(cell, fill: str) -> None:
  tc_pr = cell._tc.get_or_add_tcPr()
  shd = OxmlElement("w:shd")
  shd.set(qn("w:fill"), fill)
  tc_pr.append(shd)


def set_cell_border(cell, color: str = "D4D4D8", size: str = "6") -> None:
  tc_pr = cell._tc.get_or_add_tcPr()
  borders = tc_pr.first_child_found_in("w:tcBorders")
  if borders is None:
    borders = OxmlElement("w:tcBorders")
    tc_pr.append(borders)
  for edge in ("top", "left", "bottom", "right"):
    tag = f"w:{edge}"
    element = borders.find(qn(tag))
    if element is None:
      element = OxmlElement(tag)
      borders.append(element)
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def set_cell_padding(cell, top="90", left="120", bottom="90", right="120") -> None:
  tc_pr = cell._tc.get_or_add_tcPr()
  margin = tc_pr.first_child_found_in("w:tcMar")
  if margin is None:
    margin = OxmlElement("w:tcMar")
    tc_pr.append(margin)
  for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
    node = margin.find(qn(f"w:{side}"))
    if node is None:
      node = OxmlElement(f"w:{side}")
      margin.append(node)
    node.set(qn("w:w"), value)
    node.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
  section = doc.sections[0]
  section.top_margin = Inches(0.65)
  section.bottom_margin = Inches(0.65)
  section.left_margin = Inches(0.7)
  section.right_margin = Inches(0.7)

  styles = doc.styles
  normal = styles["Normal"]
  normal.font.name = "Aptos"
  normal.font.size = Pt(10.5)
  normal.font.color.rgb = DARK
  normal.paragraph_format.space_after = Pt(6)
  normal.paragraph_format.line_spacing = 1.08

  for name, size, color in [
    ("Title", 30, ACCENT),
    ("Heading 1", 19, ACCENT),
    ("Heading 2", 15, DARK),
    ("Heading 3", 12, DARK),
  ]:
    style = styles[name]
    style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
    style.paragraph_format.space_after = Pt(6)


def add_colored_run(paragraph, text: str, color=ACCENT, bold=True):
  run = paragraph.add_run(text)
  run.bold = bold
  run.font.color.rgb = color
  return run


def add_note(doc: Document, title: str, body: str, fill="F5F3FF", border="DDD6FE") -> None:
  table = doc.add_table(rows=1, cols=1)
  table.autofit = True
  cell = table.cell(0, 0)
  set_cell_shading(cell, fill)
  set_cell_border(cell, border, "8")
  set_cell_padding(cell, "140", "160", "140", "160")
  p = cell.paragraphs[0]
  p.paragraph_format.space_after = Pt(2)
  run = p.add_run(title)
  run.bold = True
  run.font.color.rgb = ACCENT
  run.font.size = Pt(10.5)
  p2 = cell.add_paragraph(body)
  p2.paragraph_format.space_after = Pt(0)
  for run in p2.runs:
    run.font.size = Pt(9.5)
  doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
  for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
  for item in items:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(item)


def add_image(doc: Document, filename: str, caption: str) -> None:
  path = ASSET_DIR / filename
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p.paragraph_format.keep_with_next = True
  run = p.add_run()
  run.add_picture(str(path), width=Inches(6.85))
  cap = doc.add_paragraph()
  cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
  cap.paragraph_format.space_after = Pt(10)
  r = cap.add_run(caption)
  r.italic = True
  r.font.size = Pt(9)
  r.font.color.rgb = MUTED


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
  table = doc.add_table(rows=1, cols=len(headers))
  table.style = "Table Grid"
  table.autofit = False if widths else True

  hdr = table.rows[0].cells
  for i, header in enumerate(headers):
    hdr[i].text = header
    set_cell_shading(hdr[i], "ede9fe")
    set_cell_border(hdr[i], "C4B5FD", "8")
    set_cell_padding(hdr[i])
    hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
      hdr[i].width = Inches(widths[i])
    for p in hdr[i].paragraphs:
      p.paragraph_format.space_after = Pt(0)
      for run in p.runs:
        run.bold = True
        run.font.color.rgb = DARK
        run.font.size = Pt(9.5)

  for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
      cells[i].text = value
      if widths:
        cells[i].width = Inches(widths[i])
      cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
      set_cell_border(cells[i], "E4E4E7", "4")
      set_cell_padding(cells[i])
      for p in cells[i].paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
          run.font.size = Pt(9)
          run.font.color.rgb = DARK
  doc.add_paragraph()


def add_route_table(doc: Document) -> None:
  add_table(
    doc,
    ["Area", "Key Routes", "Purpose"],
    [
      ["Public", "/, /explore, /businesses/[id], /pricing, /how-it-works, /faq, /about, /contact", "Marketing, discovery, education and public listing detail."],
      ["Auth", "/signup, /login", "Role-based registration and login."],
      ["Owner", "/dashboard/business", "Manage business listings, review requests, messages and analytics snapshot."],
      ["Investor", "/dashboard/investor", "Preferences, recommendations, saved businesses, requests and messages."],
      ["Buyer", "/dashboard/buyer", "Acquisition targets, watchlist, enquiries and owner messages."],
      ["Admin", "/admin", "Platform stats, business approval, owner verification and user management."],
      ["Shared", "/messages, /messages/[id], /profile", "Direct messaging and account profile updates."],
    ],
    [1.0, 2.55, 3.0],
  )


def build_doc() -> None:
  OUT_DIR.mkdir(parents=True, exist_ok=True)
  doc = Document()
  style_document(doc)

  # Cover
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  p.paragraph_format.space_after = Pt(0)
  r = p.add_run("Finvest")
  r.bold = True
  r.font.size = Pt(20)
  r.font.color.rgb = ACCENT

  title = doc.add_paragraph(style="Title")
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  title.add_run("Website Features and Step-by-Step User Guide")

  sub = doc.add_paragraph()
  sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
  sub.add_run("A practical manual for public visitors, business owners, investors, buyers and admins.")

  meta = doc.add_paragraph()
  meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
  meta.add_run("Prepared on May 25, 2026")

  add_note(
    doc,
    "Important positioning",
    "Finvest is a marketplace that facilitates introductions. It does not provide financial advice, handle funds, verify every financial figure, or guarantee any transaction outcome.",
    "FFF7ED",
    "FDBA74",
  )
  doc.add_page_break()

  # Contents and overview
  doc.add_heading("1. What This Website Does", level=1)
  doc.add_paragraph(
    "Finvest connects Pakistani business owners with investors and buyers. Owners list businesses, investors and buyers discover opportunities, and approved connections can message each other directly."
  )
  add_bullets(
    doc,
    [
      "Public visitors can browse public pages, view marketplace listings and read guidance pages.",
      "Business owners can create listings, submit them for admin review, boost active listings and manage connection requests.",
      "Investors can set investment preferences, browse recommendations, save businesses and send connection requests.",
      "Buyers can browse acquisition targets, save opportunities and send acquisition enquiries.",
      "Admins can approve or reject draft listings, verify owners, monitor activity and manage users.",
    ],
  )

  add_table(
    doc,
    ["Feature", "Business Owner", "Investor", "Buyer", "Admin"],
    [
      ["Create business listing", "Yes", "No", "No", "Review only"],
      ["Submit for admin review", "Yes", "No", "No", "Approve/reject"],
      ["Browse marketplace", "Yes", "Yes", "Yes", "Yes"],
      ["Save businesses", "Limited/no self-save", "Yes", "Yes", "Yes"],
      ["Send connection request", "No", "Yes", "Yes", "Admin not intended"],
      ["Accept/reject requests", "Yes, incoming", "No", "No", "Can oversee"],
      ["Direct messages", "After accepted connection", "After accepted connection", "After accepted connection", "Platform access"],
      ["User management", "No", "No", "No", "Yes"],
    ],
    [1.7, 1.25, 1.1, 1.1, 1.2],
  )

  add_image(doc, "01-marketplace.png", "UI picture: public marketplace discovery, filters and business cards.")

  doc.add_heading("2. Getting Started", level=1)
  doc.add_heading("Create an Account", level=2)
  add_numbered(
    doc,
    [
      "Open the website and choose Sign up.",
      "Select your role: Business Owner, Investor or Buyer.",
      "Enter your full name, email address, password, phone number and city.",
      "Review the legal disclaimer and create the account.",
      "After registration, the app redirects you to the correct dashboard for your role.",
    ],
  )
  doc.add_heading("Sign In", level=2)
  add_numbered(
    doc,
    [
      "Open /login.",
      "Enter your email and password.",
      "Use the password visibility icon if needed.",
      "In development, expand Demo accounts and choose a prepared role account.",
      "After login, the app redirects to /dashboard/business, /dashboard/investor, /dashboard/buyer or /admin based on role.",
    ],
  )
  add_image(doc, "02-auth.png", "UI picture: role selection, signup flow and login form.")

  doc.add_heading("3. Public Marketplace and Business Profiles", level=1)
  doc.add_heading("Explore Businesses", level=2)
  add_numbered(
    doc,
    [
      "Open /explore.",
      "Use the search bar to search by business name, industry or keyword.",
      "Open Filters to narrow results by industry, province, listing type, stage, asking price, revenue, verification and minimum trust score.",
      "Use Sort to order results by featured status, newest, trust score, asking price or revenue.",
      "Open any business card to view its full profile.",
      "Logged-in investors and buyers can save businesses from cards or profile pages.",
    ],
  )

  doc.add_heading("Business Profile Page", level=2)
  doc.add_paragraph("Each dynamic /businesses/[id] page includes the following:")
  add_bullets(
    doc,
    [
      "Cover image banner or industry visual.",
      "Business name, verified owner badge and featured badge when boosted.",
      "Industry, location, listing type and business stage.",
      "Funding required card and revenue range card.",
      "Trust score meter based on disclosure and verification signals.",
      "Full business description, highlights, business details and owner card.",
      "Connect and Save buttons.",
      "Similar businesses section.",
    ],
  )
  add_note(
    doc,
    "Connection rule",
    "Only logged-in investors and buyers can send connection requests. Business owners cannot connect to their own listings.",
  )
  add_image(doc, "03-business-profile.png", "UI picture: dynamic business profile with trust score, metrics and actions.")

  doc.add_heading("4. Business Owner Workflow", level=1)
  doc.add_paragraph("Business owners use /dashboard/business. This page is role-protected for BUSINESS_OWNER users.")
  doc.add_heading("Create and Submit a Business", level=2)
  add_numbered(
    doc,
    [
      "Sign in as a Business Owner.",
      "Open Business Dashboard and select List New Business or Add new.",
      "Enter the business name and full description.",
      "Choose industry, province, city and business stage.",
      "Enter funding required, revenue range, annual profit, employees and year established when available.",
      "Choose whether the business is for sale, seeking investment, or both.",
      "Add up to 6 key highlights.",
      "Upload up to 5 business images.",
      "Select Save Draft if you are still editing.",
      "Select Submit for Review when ready. The listing remains in DRAFT until an admin approves it.",
    ],
  )
  doc.add_heading("Manage Existing Businesses", level=2)
  add_bullets(
    doc,
    [
      "Edit updates listing details through PUT /api/businesses/:id.",
      "Delete removes the business and related saved/connection records through DELETE /api/businesses/:id.",
      "Boost toggles featured placement through POST /api/businesses/:id/boost. Only active listings can be boosted.",
      "Analytics is currently a placeholder button on listing rows.",
    ],
  )
  doc.add_heading("Manage Requests and Messages", level=2)
  add_numbered(
    doc,
    [
      "Review incoming requests in the Connection Requests table.",
      "Read the sender, business, request type, message and date.",
      "Select Accept to approve a pending request or Reject to decline it.",
      "Accepted connections can message you directly.",
      "Use Recent Messages to preview new conversations.",
      "Use the analytics snapshot chart to see recent request activity.",
    ],
  )
  add_image(doc, "04-owner-dashboard.png", "UI picture: business owner dashboard, listing form, requests and analytics snapshot.")

  doc.add_heading("5. Investor Workflow", level=1)
  doc.add_paragraph("Investors use /dashboard/investor. The dashboard focuses on recommendations and investment preferences.")
  add_numbered(
    doc,
    [
      "Sign in as an Investor.",
      "Open the Investor Dashboard.",
      "Set preferred industries, provinces and investment range.",
      "Optionally add an investment thesis, portfolio size and accreditation details.",
      "Review Recommended for You, which is shaped by your preferences and trust scores.",
      "Open a business profile to read the full listing.",
      "Save promising businesses to your watchlist.",
      "Send a connection request with a personalized message of at least 30 characters.",
      "Track request status in My Connection Requests.",
      "After an owner accepts, continue in Messages.",
    ],
  )
  add_image(doc, "05-investor-dashboard.png", "UI picture: investor preferences, recommendations, saved items and requests.")

  doc.add_heading("6. Buyer Workflow", level=1)
  doc.add_paragraph("Buyers use /dashboard/buyer. The dashboard is centered on acquisition opportunities.")
  add_numbered(
    doc,
    [
      "Sign in as a Buyer.",
      "Open the Buyer Dashboard.",
      "Review Businesses Open to Acquisition.",
      "Open a profile to inspect description, owner, metrics and trust score.",
      "Save strong targets to the Acquisition Watchlist.",
      "Send an acquisition connection request from a profile page.",
      "Track status in Acquisition Enquiries.",
      "When accepted, use Messages from Owners to continue the discussion.",
    ],
  )
  add_image(doc, "06-buyer-dashboard.png", "UI picture: buyer dashboard with acquisition targets and enquiries.")

  doc.add_heading("7. Admin Workflow", level=1)
  doc.add_paragraph("Admins use /admin. This page is role-protected for ADMIN users.")
  add_bullets(
    doc,
    [
      "Control Centre shows total users, businesses, connections, active businesses, accepted connections and pending reviews.",
      "Growth charts show user growth, listing trends and connection metrics.",
      "Pending Verifications lists DRAFT businesses awaiting review.",
      "Approve changes a listing to ACTIVE; Reject keeps or moves it out of public visibility depending on the API action.",
      "Verify Owner adds/removes an owner verification signal and affects trust score.",
      "All Users supports search, role filtering, suspend/reactivate and delete actions.",
      "Reported Content is a placeholder queue for future user reports.",
    ],
  )
  add_image(doc, "07-admin.png", "UI picture: admin control centre with stats, chart, reviews and user management.")

  doc.add_heading("8. Messages and Profile", level=1)
  doc.add_heading("Messages", level=2)
  add_bullets(
    doc,
    [
      "Messages are available to authenticated users.",
      "The inbox groups messages by conversation partner.",
      "Unread counts appear beside the page title and threads.",
      "Opening a conversation marks unread messages from that partner as read.",
      "Direct messaging should be used after accepted connection requests.",
    ],
  )
  doc.add_heading("Profile", level=2)
  add_bullets(
    doc,
    [
      "Users can update full name, phone, city, company/organisation and bio.",
      "Profile badges show role and verification status.",
      "A complete profile improves trust and gives other users more context.",
    ],
  )
  add_image(doc, "08-messages-profile.png", "UI picture: direct messages and profile edit form.")

  doc.add_heading("9. Pricing, Help and Safety Pages", level=1)
  add_bullets(
    doc,
    [
      "/pricing explains Basic, Boost and Premium plans. Basic is free; Boost and Premium increase visibility and unlock higher-tier support features.",
      "/how-it-works explains the role-based journey for owners, investors and buyers.",
      "/faq answers account, listing, verification, trust score, connections, messaging, premium, safety, legal and due diligence questions.",
      "/about and /contact provide company information and contact paths.",
    ],
  )
  add_note(
    doc,
    "Due diligence reminder",
    "All listed financial figures are self-reported. Investors and buyers should request documents, consult a qualified CA/lawyer/advisor, visit the business where appropriate and verify claims independently.",
    "FFF7ED",
    "FDBA74",
  )

  doc.add_heading("10. Technical Route and API Reference", level=1)
  add_route_table(doc)
  add_table(
    doc,
    ["API Area", "Endpoints", "Used For"],
    [
      ["Authentication", "POST /api/auth/register, POST /api/auth/login, POST /api/auth/logout, GET /api/auth/me", "Register, sign in/out and fetch the current user."],
      ["Businesses", "GET /api/businesses, POST /api/businesses, GET/PUT/DELETE /api/businesses/:id", "Marketplace search, create listings and update/delete owned listings."],
      ["Boosting", "POST /api/businesses/:id/boost", "Toggle featured placement on active owned listings."],
      ["Saving", "POST/DELETE /api/users/save/:businessId and /api/businesses/:id/save", "Save or unsave businesses for watchlists."],
      ["Connections", "POST /api/connections, GET /api/connections, PUT /api/connections/:id", "Send requests, list outgoing/incoming requests, accept/reject requests."],
      ["Messages", "GET/POST /api/messages, GET /api/messages/:userId, /api/conversations", "Inbox, direct conversations and message sending."],
      ["Investor Profile", "GET/PUT /api/investor-profile", "Store investor preferences used for recommendations."],
      ["Admin", "/api/admin/stats, /api/admin/verify/:businessId, /api/admin/trust/:businessId, /api/admin/users", "Stats, listing approval, owner verification and user management."],
      ["Profile", "PATCH /api/profile", "Update user profile fields."],
    ],
    [1.25, 3.1, 2.15],
  )

  doc.add_heading("11. Recommended Use Checklist", level=1)
  add_table(
    doc,
    ["User Type", "Before Starting", "Main Routine"],
    [
      ["Business Owner", "Prepare description, financials, photos and highlights.", "Submit listing, respond to requests, keep profile complete and use boost after approval."],
      ["Investor", "Define industries, provinces and investment budget.", "Filter/review businesses, save opportunities, send thoughtful requests and perform due diligence."],
      ["Buyer", "Know acquisition budget, preferred sectors and target geography.", "Review acquisition listings, save targets, send enquiries and continue accepted talks in messages."],
      ["Admin", "Monitor DRAFT queue and user reports.", "Approve quality listings, verify owners, manage users and watch platform metrics."],
    ],
    [1.3, 2.55, 2.75],
  )

  doc.add_paragraph()
  end = doc.add_paragraph()
  end.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run = end.add_run("End of guide")
  run.bold = True
  run.font.color.rgb = ACCENT

  doc.save(DOCX_PATH)


if __name__ == "__main__":
  build_doc()
  print(DOCX_PATH)
